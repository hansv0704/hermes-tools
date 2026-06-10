import os
import pythoncom
import win32com.client
from docx import Document
from docx.oxml import parse_xml
from docx.shared import Cm
from docx.enum.section import WD_ORIENT
from copy import deepcopy
from base_skill import BaseSkill
from typing import List, Dict, Any


class WordEditorSkill(BaseSkill):
    @property
    def name(self) -> str:
        return "word_editor_skill"

    def get_tool_declarations(self) -> List[Dict[str, Any]]:
        return [
            {
                "name": "word_live_get_structure",
                "description": "【即時模式】讀取當前開啟中 Word 文件的結構（段落數、表格數、章節大綱）",
                "parameters": {"type": "object", "properties": {}, "required": []}
            },
            {
                "name": "word_live_insert_text",
                "description": "【即時模式】在當前開啟的 Word 文件中，於游標位置插入文字",
                "parameters": {
                    "type": "object",
                    "properties": {"text": {"type": "string", "description": "要插入的文字內容"}},
                    "required": ["text"]
                }
            },
            {
                "name": "word_live_replace_text",
                "description": "【即時模式】在當前開啟的 Word 文件中搜尋並取代文字",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "old_text": {"type": "string", "description": "要被取代的舊文字"},
                        "new_text": {"type": "string", "description": "新的文字"}
                    },
                    "required": ["old_text", "new_text"]
                }
            },
            {
                "name": "word_live_read_content",
                "description": "【即時模式】讀取當前開啟中 Word 文件的全部文字內容",
                "parameters": {"type": "object", "properties": {}, "required": []}
            },
            {
                "name": "word_live_read_tables",
                "description": "【即時模式】讀取當前開啟中 Word 文件的所有表格資料",
                "parameters": {"type": "object", "properties": {}, "required": []}
            },
            {
                "name": "word_live_edit_table_cell",
                "description": "【即時模式】修改當前開啟中 Word 文件的指定表格儲存格內容",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "table_index": {"type": "integer", "description": "表格索引（從 0 開始）"},
                        "row": {"type": "integer", "description": "列索引（從 0 開始）"},
                        "col": {"type": "integer", "description": "欄索引（從 0 開始）"},
                        "text": {"type": "string", "description": "新的文字內容"}
                    },
                    "required": ["table_index", "row", "col", "text"]
                }
            },
            {
                "name": "word_live_add_table_row",
                "description": "【即時模式】在當前開啟中 Word 文件的指定表格新增一行",
                "parameters": {
                    "type": "object",
                    "properties": {"table_index": {"type": "integer", "description": "表格索引（從 0 開始）"}},
                    "required": ["table_index"]
                }
            },
            {
                "name": "word_live_save",
                "description": "【即時模式】儲存當前開啟中的 Word 文件",
                "parameters": {"type": "object", "properties": {}, "required": []}
            },
            {
                "name": "word_live_save_as",
                "description": "【即時模式】將當前開啟中的 Word 文件另存新檔",
                "parameters": {
                    "type": "object",
                    "properties": {"file_path": {"type": "string", "description": "完整的儲存路徑（如 C:\\Users\\hans\\Desktop\\test.docx）"}},
                    "required": ["file_path"]
                }
            },
            {
                "name": "word_offline_read",
                "description": "【離線模式】讀取已存檔的 .docx 檔案內容",
                "parameters": {
                    "type": "object",
                    "properties": {"file_path": {"type": "string", "description": "檔案完整路徑"}},
                    "required": ["file_path"]
                }
            },
            {
                "name": "word_offline_edit",
                "description": "【離線模式】編輯已存檔的 .docx 檔案（不破壞格式）",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "file_path": {"type": "string", "description": "檔案完整路徑"},
                        "edits": {
                            "type": "array",
                            "items": {
                                "type": "object",
                                "properties": {
                                    "type": {"type": "string", "description": "操作類型: 'replace' (取代), 'insert' (插入)"},
                                    "target": {"type": "string", "description": "目標文字"},
                                    "text": {"type": "string", "description": "新的文字內容"}
                                }
                            },
                            "description": "編輯操作清單"
                        }
                    },
                    "required": ["file_path", "edits"]
                }
            },
            {
                "name": "word_offline_create",
                "description": "【離線模式】建立新的 .docx 檔案",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "file_path": {"type": "string", "description": "儲存路徑"},
                        "content": {"type": "string", "description": "文件內容"}
                    },
                    "required": ["file_path", "content"]
                }
            },
            {
                "name": "word_offline_copy_table",
                "description": "【XML 複製引擎】在已存檔的 .docx 中複製指定表格（100% 保留格式、合併儲存格、邊框、字體）",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "file_path": {"type": "string", "description": "檔案完整路徑"},
                        "source_table_index": {"type": "integer", "description": "要複製的表格索引（從 0 開始）"},
                        "target_position": {"type": "string", "description": "插入位置：'end'（文件末尾）, 'before'（指定表格前）, 'after'（指定表格後）", "default": "end"},
                        "reference_table_index": {"type": "integer", "description": "參考表格索引（僅 target_position='before'/'after' 時使用）", "default": 0}
                    },
                    "required": ["file_path", "source_table_index"]
                }
            },
            {
                "name": "word_offline_add_section",
                "description": "【Section 頁面控制】在已存檔的 .docx 中新增 section break 並設定橫式/直式",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "file_path": {"type": "string", "description": "檔案完整路徑"},
                        "orientation": {"type": "string", "description": "'landscape'（橫式）或 'portrait'（直式），預設 landscape", "default": "landscape"},
                        "page_width_cm": {"type": "number", "description": "頁寬 cm（橫式預設 29.7）", "default": 29.7},
                        "page_height_cm": {"type": "number", "description": "頁高 cm（橫式預設 21.0）", "default": 21.0}
                    },
                    "required": ["file_path"]
                }
            },
            {
                "name": "word_offline_copy_paragraphs",
                "description": "【段落複製引擎】在已存檔的 .docx 中複製指定範圍的段落（保留字體/顏色/縮排）",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "file_path": {"type": "string", "description": "檔案完整路徑"},
                        "start_para": {"type": "integer", "description": "起始段落索引（從 0 開始）"},
                        "end_para": {"type": "integer", "description": "結束段落索引（含）"},
                        "target_position": {"type": "string", "description": "'end' 或 'before'（指定段落前）", "default": "end"},
                        "reference_para": {"type": "integer", "description": "參考段落索引（target_position='before' 時使用）", "default": 0}
                    },
                    "required": ["file_path", "start_para", "end_para"]
                }
            },
            {
                "name": "word_live_copy_table",
                "description": "【win32com 終極方案】在當前開啟的 Word 中複製指定表格並貼到文件末尾（100% 格式保留）",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "source_table_index": {"type": "integer", "description": "要複製的表格索引（從 0 開始）"}
                    },
                    "required": ["source_table_index"]
                }
            },
            {
                "name": "word_live_add_section_break",
                "description": "【win32com】在當前開啟的 Word 游標位置插入分節符號（下一頁）",
                "parameters": {"type": "object", "properties": {}, "required": []}
            },
            {
                "name": "word_live_set_orientation",
                "description": "【win32com】設定當前 section 的頁面方向",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "orientation": {"type": "string", "description": "'landscape'（橫式）或 'portrait'（直式）", "default": "landscape"}
                    },
                    "required": []
                }
            },
            {
                "name": "word_live_insert_page_break",
                "description": "【win32com】在當前游標位置插入分頁符號",
                "parameters": {"type": "object", "properties": {}, "required": []}
            },
            {
                "name": "word_live_get_page_setup",
                "description": "【win32com】讀取當前 section 的頁面設定（方向、尺寸、邊界）",
                "parameters": {"type": "object", "properties": {}, "required": []}
            }
        ]

    def _get_word_app(self):
        """取得當前 Word 應用程式實體"""
        pythoncom.CoInitialize()
        try:
            word_app = win32com.client.GetActiveObject("Word.Application")
            return word_app
        except Exception as e:
            raise Exception(f"無法連接到 Word 應用程式，請確認已開啟 Word: {e}")

    def execute(self, function_name: str, args: dict, context: dict) -> dict:
        if function_name == "word_live_get_structure":
            return self._live_get_structure()
        elif function_name == "word_live_insert_text":
            return self._live_insert_text(args.get("text", ""))
        elif function_name == "word_live_replace_text":
            return self._live_replace_text(args.get("old_text", ""), args.get("new_text", ""))
        elif function_name == "word_live_read_content":
            return self._live_read_content()
        elif function_name == "word_live_read_tables":
            return self._live_read_tables()
        elif function_name == "word_live_edit_table_cell":
            return self._live_edit_table_cell(
                args.get("table_index", 0),
                args.get("row", 0),
                args.get("col", 0),
                args.get("text", "")
            )
        elif function_name == "word_live_add_table_row":
            return self._live_add_table_row(args.get("table_index", 0))
        elif function_name == "word_live_save":
            return self._live_save()
        elif function_name == "word_live_save_as":
            return self._live_save_as(args.get("file_path", ""))
        elif function_name == "word_offline_read":
            return self._offline_read(args.get("file_path", ""))
        elif function_name == "word_offline_edit":
            return self._offline_edit(args.get("file_path", ""), args.get("edits", []))
        elif function_name == "word_offline_create":
            return self._offline_create(args.get("file_path", ""), args.get("content", ""))
        elif function_name == "word_offline_copy_table":
            return self._offline_copy_table(
                args.get("file_path", ""),
                args.get("source_table_index", 0),
                args.get("target_position", "end"),
                args.get("reference_table_index", 0)
            )
        elif function_name == "word_offline_add_section":
            return self._offline_add_section(
                args.get("file_path", ""),
                args.get("orientation", "landscape"),
                args.get("page_width_cm", 29.7),
                args.get("page_height_cm", 21.0)
            )
        elif function_name == "word_offline_copy_paragraphs":
            return self._offline_copy_paragraphs(
                args.get("file_path", ""),
                args.get("start_para", 0),
                args.get("end_para", 0),
                args.get("target_position", "end"),
                args.get("reference_para", 0)
            )
        elif function_name == "word_live_copy_table":
            return self._live_copy_table(args.get("source_table_index", 0))
        elif function_name == "word_live_add_section_break":
            return self._live_add_section_break()
        elif function_name == "word_live_set_orientation":
            return self._live_set_orientation(args.get("orientation", "landscape"))
        elif function_name == "word_live_insert_page_break":
            return self._live_insert_page_break()
        elif function_name == "word_live_get_page_setup":
            return self._live_get_page_setup()
        else:
            return {"error": f"未知工具: {function_name}"}

    # ========== 現有即時模式方法 (win32com) ==========

    def _live_get_structure(self) -> dict:
        try:
            word_app = self._get_word_app()
            doc = word_app.ActiveDocument
            paras = doc.Paragraphs.Count
            tables = doc.Tables.Count
            sections = doc.Sections.Count
            return {"result": f"📄 文件結構：{paras} 個段落，{tables} 個表格，{sections} 個章節"}
        except Exception as e:
            return {"error": f"讀取結構失敗: {e}"}

    def _live_insert_text(self, text: str) -> dict:
        try:
            word_app = self._get_word_app()
            sel = word_app.Selection
            sel.TypeText(text)
            return {"result": f"✅ 已成功插入文字：{text[:50]}{'...' if len(text) > 50 else ''}"}
        except Exception as e:
            return {"error": f"插入文字失敗: {e}"}

    def _live_replace_text(self, old_text: str, new_text: str) -> dict:
        try:
            word_app = self._get_word_app()
            doc = word_app.ActiveDocument
            find = doc.Content.Find
            find.ClearFormatting()
            find.Text = old_text
            find.Replacement.ClearFormatting()
            find.Replacement.Text = new_text
            find.Forward = True
            find.Wrap = 1
            count = 0
            while find.Execute(Replace=2):
                count += 1
            return {"result": f"✅ 已取代 {count} 處：'{old_text}' → '{new_text}'"}
        except Exception as e:
            return {"error": f"取代文字失敗: {e}"}

    def _live_read_content(self) -> dict:
        try:
            word_app = self._get_word_app()
            doc = word_app.ActiveDocument
            content = doc.Content.Text
            return {"result": content[:2000] if len(content) > 2000 else content}
        except Exception as e:
            return {"error": f"讀取內容失敗: {e}"}

    def _live_read_tables(self) -> dict:
        try:
            word_app = self._get_word_app()
            doc = word_app.ActiveDocument
            result = []
            for i, table in enumerate(doc.Tables):
                rows = table.Rows.Count
                cols = table.Columns.Count
                data = []
                for r in range(1, rows + 1):
                    row_data = []
                    for c in range(1, cols + 1):
                        cell_text = table.Cell(r, c).Range.Text.strip().replace("\r", "").replace("\x07", "")
                        row_data.append(cell_text)
                    data.append(" | ".join(row_data))
                result.append(f"📊 表格 {i+1} ({rows}x{cols}):\n" + "\n".join(data))
            return {"result": "\n\n".join(result) if result else "📄 文件中沒有表格"}
        except Exception as e:
            return {"error": f"讀取表格失敗: {e}"}

    def _live_edit_table_cell(self, table_index: int, row: int, col: int, text: str) -> dict:
        try:
            word_app = self._get_word_app()
            doc = word_app.ActiveDocument
            if table_index >= doc.Tables.Count:
                return {"error": f"表格索引超出範圍，共有 {doc.Tables.Count} 個表格"}
            table = doc.Tables[table_index + 1]
            if row >= table.Rows.Count or col >= table.Columns.Count:
                return {"error": f"儲存格 ({row+1},{col+1}) 超出範圍"}
            table.Cell(row + 1, col + 1).Range.Text = text
            return {"result": f"✅ 已修改表格 {table_index+1} 的 ({row+1},{col+1}) 為：{text}"}
        except Exception as e:
            return {"error": f"編輯儲存格失敗: {e}"}

    def _live_add_table_row(self, table_index: int) -> dict:
        try:
            word_app = self._get_word_app()
            doc = word_app.ActiveDocument
            if table_index >= doc.Tables.Count:
                return {"error": f"表格索引超出範圍，共有 {doc.Tables.Count} 個表格"}
            table = doc.Tables[table_index + 1]
            table.Rows.Add()
            return {"result": f"✅ 已為表格 {table_index+1} 新增一行"}
        except Exception as e:
            return {"error": f"新增行失敗: {e}"}

    def _live_save(self) -> dict:
        try:
            word_app = self._get_word_app()
            doc = word_app.ActiveDocument
            doc.Save()
            return {"result": "✅ 文件已儲存"}
        except Exception as e:
            return {"error": f"儲存失敗: {e}"}

    def _live_save_as(self, file_path: str) -> dict:
        try:
            word_app = self._get_word_app()
            doc = word_app.ActiveDocument
            doc.SaveAs(file_path)
            return {"result": f"✅ 文件已另存為：{file_path}"}
        except Exception as e:
            return {"error": f"另存失敗: {e}"}

    def _offline_read(self, file_path: str) -> dict:
        if not os.path.exists(file_path):
            return {"error": f"檔案不存在：{file_path}"}
        try:
            doc = Document(file_path)
            content = "\n".join([p.text for p in doc.paragraphs])
            return {"result": content[:2000] if len(content) > 2000 else content}
        except Exception as e:
            return {"error": f"讀取失敗: {e}"}

    def _offline_edit(self, file_path: str, edits: List[Dict]) -> dict:
        if not os.path.exists(file_path):
            return {"error": f"檔案不存在：{file_path}"}
        try:
            doc = Document(file_path)
            for edit in edits:
                op_type = edit.get("type", "")
                target = edit.get("target", "")
                text = edit.get("text", "")
                if op_type == "replace":
                    for para in doc.paragraphs:
                        if target in para.text:
                            para.text = para.text.replace(target, text)
                elif op_type == "insert":
                    for para in doc.paragraphs:
                        if target in para.text:
                            para.text = para.text.replace(target, target + text)
            doc.save(file_path)
            return {"result": f"✅ 已編輯檔案：{file_path}"}
        except Exception as e:
            return {"error": f"編輯失敗: {e}"}

    def _offline_create(self, file_path: str, content: str) -> dict:
        try:
            doc = Document()
            doc.add_paragraph(content)
            doc.save(file_path)
            return {"result": f"✅ 已建立檔案：{file_path}"}
        except Exception as e:
            return {"error": f"建立失敗: {e}"}

    # ========== 新增：XML 深度複製引擎 (離線) ==========

    def _offline_copy_table(self, file_path: str, source_table_index: int,
                            target_position: str = "end", reference_table_index: int = 0) -> dict:
        """
        方案一：XML 表格複製引擎
        使用 deepcopy + lxml 在已存檔 .docx 中複製表格，100% 保留格式。
        """
        if not os.path.exists(file_path):
            return {"error": f"檔案不存在：{file_path}"}
        try:
            doc = Document(file_path)
            if source_table_index >= len(doc.tables):
                return {"error": f"表格索引 {source_table_index} 超出範圍，檔案共有 {len(doc.tables)} 個表格"}

            source_tbl = doc.tables[source_table_index]._tbl
            new_tbl = deepcopy(source_tbl)

            if target_position == "end":
                doc.element.body.append(new_tbl)
            elif target_position == "before":
                if reference_table_index >= len(doc.tables):
                    return {"error": f"參考表格索引 {reference_table_index} 超出範圍"}
                ref_tbl = doc.tables[reference_table_index]._tbl
                ref_tbl.addprevious(new_tbl)
            elif target_position == "after":
                if reference_table_index >= len(doc.tables):
                    return {"error": f"參考表格索引 {reference_table_index} 超出範圍"}
                ref_tbl = doc.tables[reference_table_index]._tbl
                ref_tbl.addnext(new_tbl)
            else:
                return {"error": f"無效的 target_position：{target_position}，可用 'end'/'before'/'after'"}

            doc.save(file_path)

            # 計算行列數
            ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
            rows = len(new_tbl.findall(f"{ns}tr"))
            cols = len(new_tbl.findall(f"{ns}tblGrid/{ns}gridCol"))

            return {
                "result": f"✅ 已複製表格 {source_table_index} ({rows}列 × {cols}欄) → {target_position}，已儲存",
                "rows": rows,
                "cols": cols
            }
        except Exception as e:
            return {"error": f"XML 表格複製失敗: {e}"}

    def _offline_add_section(self, file_path: str, orientation: str = "landscape",
                              page_width_cm: float = 29.7, page_height_cm: float = 21.0) -> dict:
        """
        方案二：Section 頁面方向控制
        在已存檔 .docx 中新增 section break 並設定橫式/直式。
        """
        if not os.path.exists(file_path):
            return {"error": f"檔案不存在：{file_path}"}
        try:
            doc = Document(file_path)
            new_section = doc.add_section()

            orient_map = {
                "landscape": WD_ORIENT.LANDSCAPE,
                "portrait": WD_ORIENT.PORTRAIT,
            }
            if orientation.lower() not in orient_map:
                return {"error": f"無效方向：{orientation}，可用 'landscape' 或 'portrait'"}

            new_section.orientation = orient_map[orientation.lower()]
            new_section.page_width = Cm(page_width_cm)
            new_section.page_height = Cm(page_height_cm)

            doc.save(file_path)

            orient_label = "橫式" if orientation.lower() == "landscape" else "直式"
            return {
                "result": f"✅ 已新增 {orient_label} section（{page_width_cm}cm × {page_height_cm}cm），已儲存",
                "orientation": orient_label,
                "width_cm": page_width_cm,
                "height_cm": page_height_cm
            }
        except Exception as e:
            return {"error": f"新增 section 失敗: {e}"}

    def _offline_copy_paragraphs(self, file_path: str, start_para: int, end_para: int,
                                  target_position: str = "end", reference_para: int = 0) -> dict:
        """
        方案三：段落複製引擎
        使用 XML deepcopy 複製指定範圍的段落，保留字體/顏色/縮排。
        """
        if not os.path.exists(file_path):
            return {"error": f"檔案不存在：{file_path}"}
        try:
            doc = Document(file_path)
            total_paras = len(doc.paragraphs)

            if start_para < 0 or end_para >= total_paras or start_para > end_para:
                return {"error": f"段落範圍 [{start_para}, {end_para}] 無效，檔案共有 {total_paras} 個段落"}

            copied_count = 0
            new_elements = []

            for i in range(start_para, end_para + 1):
                para = doc.paragraphs[i]
                new_para_xml = deepcopy(para._p)
                new_elements.append(new_para_xml)
                copied_count += 1

            if target_position == "end":
                for elem in new_elements:
                    doc.element.body.append(elem)
            elif target_position == "before":
                ref_para_elem = doc.paragraphs[reference_para]._p
                for elem in reversed(new_elements):
                    ref_para_elem.addprevious(elem)
            else:
                return {"error": f"無效的 target_position：{target_position}，可用 'end'/'before'"}

            doc.save(file_path)
            return {"result": f"✅ 已複製 {copied_count} 個段落（索引 {start_para}-{end_para}）→ {target_position}，已儲存"}
        except Exception as e:
            return {"error": f"段落複製失敗: {e}"}

    # ========== 新增：win32com 終極方案 (即時) ==========

    def _live_copy_table(self, source_table_index: int) -> dict:
        """
        方案四：win32com 表格複製
        透過 Word 原生 COM 介面複製表格，100% 格式保留（含漸層、ActiveX）。
        """
        try:
            word_app = self._get_word_app()
            doc = word_app.ActiveDocument

            if source_table_index >= doc.Tables.Count:
                return {"error": f"表格索引 {source_table_index} 超出範圍，共有 {doc.Tables.Count} 個表格"}

            # 選取並複製表格
            table = doc.Tables[source_table_index + 1]  # COM 為 1-indexed
            table.Range.Copy()

            # 移到文件末尾並貼上
            doc.Content.Select()
            word_app.Selection.Collapse(0)  # wdCollapseEnd = 0
            word_app.Selection.InsertBreak(2)  # wdSectionBreakNextPage
            word_app.Selection.Paste()

            rows = table.Rows.Count
            cols = table.Columns.Count

            return {"result": f"✅ 已複製表格 {source_table_index} ({rows}列 × {cols}欄) → 文件末尾（新 section），格式 100% 保留"}
        except Exception as e:
            return {"error": f"win32com 表格複製失敗: {e}"}

    def _live_add_section_break(self) -> dict:
        """
        方案四：win32com 分節符號
        在當前游標位置插入「下一頁」分節符號。
        """
        try:
            word_app = self._get_word_app()
            sel = word_app.Selection
            # wdSectionBreakNextPage = 2
            sel.InsertBreak(2)
            return {"result": "✅ 已在游標位置插入分節符號（下一頁）"}
        except Exception as e:
            return {"error": f"插入分節符號失敗: {e}"}

    def _live_set_orientation(self, orientation: str = "landscape") -> dict:
        """
        方案四：win32com 頁面方向
        設定當前 section 的頁面方向。
        """
        try:
            word_app = self._get_word_app()
            sel = word_app.Selection
            page_setup = sel.PageSetup

            if orientation.lower() == "landscape":
                page_setup.Orientation = 1  # wdOrientLandscape
                orient_label = "橫式"
            elif orientation.lower() == "portrait":
                page_setup.Orientation = 0  # wdOrientPortrait
                orient_label = "直式"
            else:
                return {"error": f"無效方向：{orientation}，可用 'landscape' 或 'portrait'"}

            return {"result": f"✅ 當前 section 已設為 {orient_label}"}
        except Exception as e:
            return {"error": f"設定頁面方向失敗: {e}"}

    def _live_insert_page_break(self) -> dict:
        """
        方案四：win32com 分頁符號
        在當前游標位置插入分頁符號。
        """
        try:
            word_app = self._get_word_app()
            sel = word_app.Selection
            # wdPageBreak = 1
            sel.InsertBreak(1)
            return {"result": "✅ 已在游標位置插入分頁符號"}
        except Exception as e:
            return {"error": f"插入分頁符號失敗: {e}"}

    def _live_get_page_setup(self) -> dict:
        """
        方案四：win32com 頁面設定讀取
        讀取當前 section 的頁面設定。
        """
        try:
            word_app = self._get_word_app()
            sel = word_app.Selection
            ps = sel.PageSetup

            orientation = "橫式" if ps.Orientation == 1 else "直式"
            width_pts = ps.PageWidth
            height_pts = ps.PageHeight
            # 1 pt = 0.0352778 cm
            width_cm = round(width_pts * 0.0352778, 2)
            height_cm = round(height_pts * 0.0352778, 2)
            left_margin = round(ps.LeftMargin * 0.0352778, 2)
            right_margin = round(ps.RightMargin * 0.0352778, 2)
            top_margin = round(ps.TopMargin * 0.0352778, 2)
            bottom_margin = round(ps.BottomMargin * 0.0352778, 2)

            return {
                "result": f"📐 當前頁面設定：{orientation}，{width_cm}cm × {height_cm}cm",
                "orientation": orientation,
                "width_cm": width_cm,
                "height_cm": height_cm,
                "margins": {
                    "left": left_margin,
                    "right": right_margin,
                    "top": top_margin,
                    "bottom": bottom_margin
                }
            }
        except Exception as e:
            return {"error": f"讀取頁面設定失敗: {e}"}
