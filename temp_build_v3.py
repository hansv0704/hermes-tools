"""
重建 表1.1-2_拆分版_v3.docx
表 b 9欄：加回「優先辦理區編號」和「大規模崩塌潛勢區」
"""
from docx import Document
from docx.shared import Cm, Emu, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

SRC = r"C:\Users\hans\Desktop\Alice_Brain_Arch_20260506_031953\作業區\表1.1-2_拆分版_v2.docx"
DST = r"C:\Users\hans\Desktop\Alice_Brain_Arch_20260506_031953\作業區\表1.1-2_拆分版_v3.docx"

def set_cell_format(cell, text, font_name=None, font_size=Pt(14), is_header=False, bg_color=None):
    """設定儲存格格式：清除段落、加入文字、字體、行距、垂直置中、邊框、底色"""
    # 清除現有段落
    for p in cell.paragraphs:
        p.clear()
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing = Pt(20)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    
    run = p.add_run(str(text))
    run.font.size = font_size
    if is_header and font_name:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    # data cells: 不設字體 = 繼承 Normal
    
    # 垂直置中
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # 邊框
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_pr.append(parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'</w:tcBorders>'))
    
    # 底色
    if bg_color:
        tc_pr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}" w:val="clear"/>'))


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = '標楷體'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
    run.font.size = Pt(14)


def main():
    src = Document(SRC)
    dst = Document()
    
    # 複製 section
    for sec in src.sections:
        ns = dst.sections[0]
        ns.orientation = sec.orientation
        ns.page_width = sec.page_width
        ns.page_height = sec.page_height
        ns.left_margin = sec.left_margin
        ns.right_margin = sec.right_margin
        ns.top_margin = sec.top_margin
        ns.bottom_margin = sec.bottom_margin
    
    ta_src = src.tables[0]
    tb_src = src.tables[1]
    nrows = len(ta_src.rows)  # 6
    
    # ========== 表 a (8欄，直接複製) ==========
    add_title(dst, '表 1.1-2a\u3000各計畫區基本資料與地質特性一覽表')
    
    ta = dst.add_table(rows=nrows, cols=8)
    ta.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 欄寬：從 src 取得，或等分
    a_col_widths = []
    for ci in range(8):
        w = ta_src.columns[ci].width
        if w is None:
            w = 1107916  # 8863330/8
        a_col_widths.append(w)
        ta.columns[ci].width = w
    
    for j in range(nrows):
        for ci in range(8):
            tc = ta.cell(j, ci)
            tc._tc.get_or_add_tcPr().append(
                parse_xml(f'<w:tcW {nsdecls("w")} w:w="{a_col_widths[ci]}" w:type="dxa"/>'))
            text = ta_src.cell(j, ci).text
            set_cell_format(tc, text, font_name='標楷體', is_header=(j==0),
                          bg_color='F2F2F2' if j==0 else None)
    
    # ========== 換頁 ==========
    dst.add_page_break()
    
    # ========== 表 b (9欄) ==========
    add_title(dst, '表 1.1-2b\u3000各計畫區不連續面分析與邊坡活動性一覽表')
    
    # 欄寬 (總 8863330 EMU)
    CW = [483330, 800000, 1480000, 1300000, 900000, 1100000, 900000, 700000, 1200000]
    HDR_B = ['編號', '優先辦理區\n編號', '大規模崩塌\n潛勢區', '不連續面位態',
             '不連續面與\n坡面關係', '主要滑動機制', '滑動面深度\n(m) 地表下',
             '地下水位\n(m) 地表下', '邊坡活動性']
    
    tb = dst.add_table(rows=nrows, cols=9)
    tb.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, w in enumerate(CW):
        tb.columns[ci].width = w
    
    for j in range(nrows):
        if j == 0:
            row_data = HDR_B
        else:
            row_data = [
                ta_src.cell(j, 0).text.strip(),   # 編號
                ta_src.cell(j, 1).text.strip(),   # 優先區
                ta_src.cell(j, 2).text.strip(),   # 潛勢區
                tb_src.cell(j, 1).text.strip(),   # 位態
                tb_src.cell(j, 2).text.strip(),   # 坡面關係
                tb_src.cell(j, 3).text.strip(),   # 滑動機制
                tb_src.cell(j, 4).text.strip(),   # 深度
                tb_src.cell(j, 5).text.strip(),   # 水位
                tb_src.cell(j, 6).text.strip(),   # 活動性
            ]
        for ci in range(9):
            tc = tb.cell(j, ci)
            tc._tc.get_or_add_tcPr().append(
                parse_xml(f'<w:tcW {nsdecls("w")} w:w="{CW[ci]}" w:type="dxa"/>'))
            set_cell_format(tc, row_data[ci], font_name='標楷體', is_header=(j==0),
                          bg_color='F2F2F2' if j==0 else None)
    
    dst.save(DST)
    print(f"✅ v3 已儲存: {DST}")
    print(f"   表a: {nrows}列 x 8欄, 總寬={sum(a_col_widths)} EMU = {round(sum(a_col_widths)/360000,2)} cm")
    print(f"   表b: {nrows}列 x 9欄, 總寬={sum(CW)} EMU = {round(sum(CW)/360000,2)} cm")
    print(f"   可用寬度: 8863330 EMU = 24.62 cm")
    print(f"   Header: 標楷體 14pt, 灰底 #F2F2F2")
    print(f"   Data: 繼承 Normal (中文標楷體/英文TNR), 14pt, 固定行距20pt")

if __name__ == '__main__':
    main()
