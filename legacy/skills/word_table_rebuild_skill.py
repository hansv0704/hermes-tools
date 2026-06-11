
"""
word_table_rebuild_skill.py — Word 表格重建專用 Skill
封裝 v2→v5 成功流程：從原始表格擷取格式 → 參數化重建 → 驗證輸出

核心經驗（來自 6 次試錯）：
  1. 新建獨立 .docx，不修改原始檔
  2. Data cells 不設字體 → 繼承 Normal（中文=標楷體, 英文=TNR）
  3. Header 設定 標楷體 14pt + #F2F2F2 灰底
  4. 行距固定 20pt（14pt字÷0.7）
  5. 欄寬從原始表格等比例縮放
  6. 寫入後強制驗證
"""
import os
from docx import Document
from docx.shared import Cm, Pt, Emu, Twips
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from lxml import etree
from copy import deepcopy
from base_skill import BaseSkill
from typing import List, Dict, Any, Optional

# OOXML namespace
WML_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


class WordTableRebuildSkill(BaseSkill):
    """Word 表格重建引擎 — 從原始檔擷取格式後參數化重建"""

    @property
    def name(self) -> str:
        return "word_table_rebuild_skill"

    def get_tool_declarations(self) -> List[Dict[str, Any]]:
        return [
            {
                "name": "word_table_rebuild",
                "description": "【Word 表格重建】從原始檔案擷取表格格式並重建為多個拆分表格。封裝 v2→v5 成功 9 步流程。",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "source_path": {
                            "type": "string",
                            "description": "原始 Word 檔案路徑（用於擷取表格格式）"
                        },
                        "source_table_index": {
                            "type": "integer",
                            "description": "原始檔案中要擷取格式的表格索引（從 0 開始）"
                        },
                        "split_definitions": {
                            "type": "array",
                            "items": {
                                "type": "object",
                                "properties": {
                                    "columns": {
                                        "type": "array",
                                        "items": {"type": "string"},
                                        "description": "此拆分表的欄位名稱清單"
                                    },
                                    "source_col_indices": {
                                        "type": "array",
                                        "items": {"type": "integer"},
                                        "description": "對應原始表格的欄索引（從 0 開始）"
                                    },
                                    "title": {"type": "string", "description": "表格標題（選填）"}
                                }
                            },
                            "description": "拆分定義清單"
                        },
                        "output_dir": {
                            "type": "string",
                            "description": "輸出路徑（完整 .docx 路徑）"
                        },
                        "source_section_index": {
                            "type": "integer",
                            "description": "原始檔案中要複製頁面設定的 section 索引（從 0 開始）",
                            "default": 0
                        }
                    },
                    "required": ["source_path", "source_table_index", "split_definitions", "output_dir"]
                }
            },
            {
                "name": "word_get_stored_paths",
                "description": "【路徑查詢】從 DuckDB 查詢已記錄的 Word 檔案路徑，避免每次重新搜尋。",
                "parameters": {
                    "type": "object",
                    "properties": {},
                    "required": []
                }
            }
        ]

    def execute(self, function_name: str, args: dict, context: dict) -> dict:
        if function_name == "word_table_rebuild":
            return self._rebuild(
                source_path=args.get("source_path", ""),
                source_table_index=args.get("source_table_index", 0),
                split_definitions=args.get("split_definitions", []),
                output_dir=args.get("output_dir", ""),
                source_section_index=args.get("source_section_index", 0),
            )
        elif function_name == "word_get_stored_paths":
            return self._get_stored_paths()
        else:
            return {"error": f"未知工具: {function_name}"}

    # ─── 主重建方法 ───

    def _rebuild(self, source_path: str, source_table_index: int,
                 split_definitions: list, output_dir: str,
                 source_section_index: int = 0) -> dict:
        """從原始表格擷取格式並重建拆分表"""
        if not os.path.exists(source_path):
            return {"error": f"原始檔案不存在：{source_path}"}

        try:
            src_doc = Document(source_path)

            # ── Step 1: 從原始 section 複製頁面設定 ──
            page_config = self._extract_section_config(src_doc, source_section_index)

            # ── Step 2: 從原始表格擷取欄寬 ──
            if source_table_index >= len(src_doc.tables):
                return {"error": f"表格索引 {source_table_index} 超出範圍，共 {len(src_doc.tables)} 個表格"}
            src_table = src_doc.tables[source_table_index]
            src_grid = self._extract_grid_col_widths(src_table)
            if not src_grid:
                return {"error": "無法擷取原始表格欄寬"}

            # ── Step 3: 擷取表格內文（所有 rows） ──
            table_data = self._extract_table_data(src_table)

            # ── Step 4: 計算可用寬度 ──
            usable_width_cm = page_config["page_width_cm"] - page_config["margin_left_cm"] - page_config["margin_right_cm"]

            # ── Step 5: 建立輸出文件 ──
            out_doc = Document()
            self._apply_section_config(out_doc, page_config)

            report = {"tables": []}

            for defn in split_definitions:
                cols = defn.get("columns", [])
                src_indices = defn.get("source_col_indices", [])
                title = defn.get("title", "")

                # 等比例計算欄寬
                raw_widths = [src_grid[i] for i in src_indices if i < len(src_grid)]
                total_raw = sum(raw_widths) if raw_widths else 1
                usable_emu = int(usable_width_cm * 360000)
                scaled_widths = [int((w / total_raw) * usable_emu) for w in raw_widths]

                # 建立表格
                n_cols = len(cols)
                n_rows = len(table_data) + 1  # +1 for header
                table = out_doc.add_table(rows=n_rows, cols=n_cols)
                table.style = 'Table Grid'

                # 設定欄寬
                tbl = table._tbl
                tblGrid = tbl.find(qn('w:tblGrid'))
                if tblGrid is not None:
                    for gc in tblGrid.findall(qn('w:gridCol')):
                        tblGrid.remove(gc)
                    for w in scaled_widths:
                        gc = etree.SubElement(tblGrid, qn('w:gridCol'))
                        gc.set(qn('w:w'), str(w))

                # ── Header row ──
                for c_idx, col_name in enumerate(cols):
                    cell = table.cell(0, c_idx)
                    self._set_cell_shading(cell, "F2F2F2")
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    self._set_paragraph_spacing(p, line_spacing_pt=20)
                    run = p.runs[0] if p.runs else p.add_run(col_name)
                    if not p.runs:
                        p.add_run(col_name)
                    run = p.runs[0]
                    run.font.size = Pt(14)
                    run.font.bold = True
                    # Header: 中英文都用標楷體
                    rPr = run._r.get_or_add_rPr()
                    rFonts = rPr.find(qn('w:rFonts'))
                    if rFonts is None:
                        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="標楷體" w:hAnsi="標楷體" w:eastAsia="標楷體"/>')
                        rPr.insert(0, rFonts)
                    else:
                        rFonts.set(qn('w:ascii'), '標楷體')
                        rFonts.set(qn('w:hAnsi'), '標楷體')
                        rFonts.set(qn('w:eastAsia'), '標楷體')

                # ── Data rows ──
                for r_idx, row_data in enumerate(table_data):
                    for c_idx, src_col in enumerate(src_indices):
                        if c_idx >= n_cols:
                            break
                        if src_col < len(row_data):
                            cell = table.cell(r_idx + 1, c_idx)
                            p = cell.paragraphs[0]
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            self._set_paragraph_spacing(p, line_spacing_pt=20)
                            text = row_data[src_col]
                            run = p.runs[0] if p.runs else p.add_run(text)
                            if not p.runs:
                                p.add_run(text)
                            run = p.runs[0]
                            run.font.size = Pt(14)
                            # 🔑 Data cells: 不設定字體 → 繼承 Normal（中文=標楷體, 英文=TNR）
                            # 這是 v2 成功的關鍵！

                report["tables"].append({
                    "title": title,
                    "columns": n_cols,
                    "rows": n_rows,
                    "widths_twips": [round(w / 635) for w in scaled_widths]  # EMU → twips 近似
                })

                if title:
                    out_doc.add_paragraph()

            # ── Step 6: 儲存 ──
            out_doc.save(output_dir)

            # ── Step 7: 驗證 ──
            verify_result = self._verify_output(output_dir, len(split_definitions), page_config)

            return {
                "result": f"✅ 已重建 {len(split_definitions)} 個表格 → {output_dir}",
                "page_config": page_config,
                "tables": report["tables"],
                "verification": verify_result
            }

        except Exception as e:
            return {"error": f"表格重建失敗: {e}"}

    # ─── 輔助方法 ───

    def _extract_section_config(self, doc: Document, section_index: int) -> dict:
        """擷取 section 的頁面設定"""
        if section_index >= len(doc.sections):
            section_index = 0
        sec = doc.sections[section_index]
        return {
            "page_width_cm": sec.page_width.cm,
            "page_height_cm": sec.page_height.cm,
            "orientation": "landscape" if sec.orientation == WD_ORIENT.LANDSCAPE else "portrait",
            "margin_left_cm": sec.left_margin.cm,
            "margin_right_cm": sec.right_margin.cm,
            "margin_top_cm": sec.top_margin.cm,
            "margin_bottom_cm": sec.bottom_margin.cm,
        }

    def _extract_grid_col_widths(self, table) -> list:
        """從表格 XML 擷取 gridCol 寬度清單（EMU）"""
        tbl = table._tbl
        grid = tbl.find(qn('w:tblGrid'))
        if grid is None:
            return []
        widths = []
        for gc in grid.findall(qn('w:gridCol')):
            w = gc.get(qn('w:w'))
            widths.append(int(w) if w else 0)
        return widths

    def _extract_table_data(self, table) -> list:
        """擷取表格所有行的文字資料（不含 header）"""
        rows_data = []
        for r_idx, row in enumerate(table.rows):
            if r_idx == 0:
                continue  # skip header
            row_data = []
            for cell in row.cells:
                text = cell.text.strip().replace('\r', '').replace('\x07', '')
                row_data.append(text)
            rows_data.append(row_data)
        return rows_data

    def _apply_section_config(self, doc: Document, config: dict):
        """套用頁面設定到文件的預設 section"""
        if doc.sections:
            sec = doc.sections[0]
        else:
            sec = doc.add_section()
        sec.page_width = Cm(config["page_width_cm"])
        sec.page_height = Cm(config["page_height_cm"])
        sec.left_margin = Cm(config["margin_left_cm"])
        sec.right_margin = Cm(config["margin_right_cm"])
        sec.top_margin = Cm(config["margin_top_cm"])
        sec.bottom_margin = Cm(config["margin_bottom_cm"])
        if config.get("orientation") == "landscape":
            sec.orientation = WD_ORIENT.LANDSCAPE

    def _set_cell_shading(self, cell, hex_color: str):
        """設定儲存格底色"""
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>'
        )
        cell._tc.get_or_add_tcPr().append(shading)

    def _set_paragraph_spacing(self, para, line_spacing_pt: float = 20):
        """設定段落行距為固定值"""
        pPr = para._p.get_or_add_pPr()
        spacing = pPr.find(qn('w:spacing'))
        if spacing is None:
            spacing = parse_xml(f'<w:spacing {nsdecls("w")}/>')
            pPr.append(spacing)
        # 固定行距：20pt = 400 twips (1pt = 20 twips)
        spacing.set(qn('w:line'), str(int(line_spacing_pt * 20)))
        spacing.set(qn('w:lineRule'), 'exact')

    def _verify_output(self, file_path: str, expected_table_count: int, page_config: dict) -> dict:
        """驗證輸出檔案"""
        issues = []
        try:
            doc = Document(file_path)
            actual_tables = len(doc.tables)
            if actual_tables != expected_table_count:
                issues.append(f"表格數: 預期{expected_table_count}, 實際{actual_tables}")

            # 檢查 section 設定
            if doc.sections:
                sec = doc.sections[0]
                w_diff = abs(sec.page_width.cm - page_config["page_width_cm"])
                if w_diff > 0.5:
                    issues.append(f"頁寬偏差: {w_diff:.1f}cm")

            status = "pass" if not issues else "warn"
            return {"status": status, "issues": issues, "table_count": actual_tables}
        except Exception as e:
            return {"status": "fail", "issues": [str(e)]}

    def _get_stored_paths(self) -> dict:
        """從 DuckDB 查詢已記錄的 Word 檔案路徑（此方法需在執行時由 agent 調用 manage_data_hub）"""
        return {
            "result": "請使用 manage_data_hub 查詢: SELECT * FROM system_facts WHERE fact_key LIKE 'word_file:%'"
        }
